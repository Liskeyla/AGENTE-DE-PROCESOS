import os
import re
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument
import openpyxl

from app.core.config import settings


class DocumentProcessor:
    """Extrae texto de múltiples formatos de archivo."""

    SUPPORTED_EXTENSIONS = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".xlsx": "xlsx",
        ".xls": "xlsx",
        ".txt": "txt",
        ".csv": "csv",
    }

    @staticmethod
    def detect_file_type(filename: str) -> Optional[str]:
        ext = Path(filename).suffix.lower()
        return DocumentProcessor.SUPPORTED_EXTENSIONS.get(ext)

    @staticmethod
    def extract_text(file_path: str, file_type: str) -> str:
        extractors = {
            "pdf": DocumentProcessor._extract_pdf,
            "docx": DocumentProcessor._extract_docx,
            "xlsx": DocumentProcessor._extract_xlsx,
            "txt": DocumentProcessor._extract_txt,
            "csv": DocumentProcessor._extract_csv,
        }
        extractor = extractors.get(file_type)
        if not extractor:
            raise ValueError(f"Tipo de archivo no soportado: {file_type}")
        return extractor(file_path)

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        doc = fitz.open(file_path)
        pages = []
        for page in doc:
            pages.append(page.get_text())
        doc.close()
        return "\n\n".join(pages)

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)

    @staticmethod
    def _extract_xlsx(file_path: str) -> str:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"--- Hoja: {sheet_name} ---")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append(" | ".join(cells))
        wb.close()
        return "\n".join(parts)

    @staticmethod
    def _extract_txt(file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    @staticmethod
    def _extract_csv(file_path: str) -> str:
        import csv
        rows = []
        with open(file_path, "r", encoding="utf-8", errors="replace", newline="") as f:
            reader = csv.reader(f)
            for row in reader:
                if any(cell.strip() for cell in row):
                    rows.append(" | ".join(row))
        return "\n".join(rows)

    @staticmethod
    def clean_text(text: str) -> str:
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"^\s+|\s+$", "", text, flags=re.MULTILINE)
        return text.strip()

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1500, overlap: int = 150) -> list[dict]:
        if len(text) <= chunk_size:
            return [{"index": 0, "content": text}]

        chunks = []
        start = 0
        index = 0
        while start < len(text):
            end = start + chunk_size
            if end < len(text):
                break_point = text.rfind("\n", start, end)
                if break_point > start:
                    end = break_point
            chunks.append({"index": index, "content": text[start:end].strip()})
            start = end - overlap
            index += 1
        return chunks

    @staticmethod
    def ensure_upload_dir(project_id: str) -> str:
        upload_path = os.path.join(settings.UPLOAD_DIR, str(project_id))
        os.makedirs(upload_path, exist_ok=True)
        return upload_path
