"""Exporta todos los documentos SGQ de un proyecto a un único Word editable."""

from __future__ import annotations

import io
import re
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.services.sgq_document_catalog import DOCUMENT_TITLES, PROGRESSIVE_DOC_TYPES

DIAGRAM_TYPES = {"mapa_procesos", "diagrama_flujo", "organigrama"}

FIELD_LABELS = {
    "summary": "Resumen",
    "policy_text": "Política",
    "commitments": "Compromisos",
    "alignment_with_context": "Alineación con el contexto",
    "alignment_with_policy": "Alineación con la política",
    "scope_statement": "Declaración de alcance",
    "products_services": "Productos y servicios",
    "locations": "Ubicaciones",
    "boundaries": "Límites",
    "applicability_notes": "Notas de aplicabilidad",
    "exclusions": "Exclusiones",
    "monitoring_review": "Revisión / monitoreo",
    "internal_context": "Contexto interno",
    "external_context": "Contexto externo",
    "stakeholders": "Partes interesadas",
    "characterizations": "Caracterizaciones",
    "interactions": "Interacciones",
    "requirements": "Requisitos",
    "objectives": "Objetivos",
    "procedures": "Procedimientos",
    "indicators": "Indicadores",
    "entries": "Entradas",
    "records": "Registros",
    "organization_name": "Organización",
    "processes": "Procesos",
    "diagrams": "Diagramas",
    "nodes": "Nodos",
}


def _as_str(value: Any, fallback: str = "") -> str:
    if value is None:
        return fallback
    if isinstance(value, bool):
        return "Sí" if value else "No"
    if isinstance(value, (int, float)):
        return str(value)
    text = str(value).strip()
    return text or fallback


def _as_list(value: Any) -> list:
    return value if isinstance(value, list) else []


def _label(key: str) -> str:
    if key in FIELD_LABELS:
        return FIELD_LABELS[key]
    return key.replace("_", " ").strip().capitalize()


def _set_run_font(run, *, bold: bool = False, size: int = 11, color: RGBColor | None = None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if color is not None:
        run.font.color.rgb = color


def _add_heading(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        _set_run_font(run, bold=True, size=16 if level == 1 else 13)
    return heading


def _add_paragraph(doc: Document, text: str, *, bold: bool = False, italic: bool = False):
    if not text:
        return None
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, bold=bold, size=11)
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def _add_bullets(doc: Document, items: list[Any]):
    for item in items:
        if isinstance(item, dict):
            text = "; ".join(
                f"{_label(k)}: {_as_str(v)}"
                for k, v in item.items()
                if v not in (None, "", [], {})
            )
        else:
            text = _as_str(item)
        if text:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(text)
            _set_run_font(run, size=11)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    if not rows:
        _add_paragraph(doc, "Sin datos registrados aún.", italic=True)
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                _set_run_font(run, bold=True, size=10)
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, value in enumerate(row):
            cells[c_idx].text = value
            for paragraph in cells[c_idx].paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, size=10)
    doc.add_paragraph()


def _join_list(value: Any) -> str:
    if isinstance(value, list):
        return ", ".join(_as_str(v) for v in value if _as_str(v))
    return _as_str(value)


def _render_contexto(doc: Document, content: dict):
    if content.get("summary"):
        _add_paragraph(doc, _as_str(content["summary"]))
    for section_key, title in (
        ("internal_context", "Contexto interno"),
        ("external_context", "Contexto externo"),
    ):
        items = _as_list(content.get(section_key))
        if not items:
            continue
        _add_heading(doc, title, level=2)
        rows = [
            [
                _as_str(item.get("factor")),
                _as_str(item.get("description")),
                _as_str(item.get("impact")),
            ]
            for item in items
            if isinstance(item, dict)
        ]
        _add_table(doc, ["Factor", "Descripción", "Impacto"], rows)
    if content.get("monitoring_review"):
        _add_heading(doc, "Revisión / monitoreo", level=2)
        _add_paragraph(doc, _as_str(content["monitoring_review"]))


def _render_alcance(doc: Document, content: dict):
    _add_paragraph(doc, _as_str(content.get("scope_statement"), "Alcance del SGC en elaboración."))
    for key, title in (
        ("products_services", "Productos y servicios"),
        ("locations", "Ubicaciones"),
    ):
        items = _as_list(content.get(key))
        if items:
            _add_heading(doc, title, level=2)
            _add_bullets(doc, items)
    if content.get("boundaries"):
        _add_heading(doc, "Límites", level=2)
        _add_paragraph(doc, _as_str(content["boundaries"]))
    if content.get("applicability_notes"):
        _add_heading(doc, "Aplicabilidad", level=2)
        _add_paragraph(doc, _as_str(content["applicability_notes"]))
    exclusions = [e for e in _as_list(content.get("exclusions")) if isinstance(e, dict)]
    if exclusions:
        _add_heading(doc, "Exclusiones", level=2)
        rows = [
            [_as_str(e.get("clause")), _as_str(e.get("justification"))]
            for e in exclusions
        ]
        _add_table(doc, ["Cláusula", "Justificación"], rows)


def _render_partes(doc: Document, content: dict):
    rows = []
    for s in _as_list(content.get("stakeholders")):
        if not isinstance(s, dict):
            continue
        rows.append(
            [
                _as_str(s.get("name")),
                _as_str(s.get("type")),
                _join_list(s.get("needs")),
                _join_list(s.get("expectations")),
                _as_str(s.get("monitoring_method")),
            ]
        )
    _add_table(
        doc,
        ["Parte interesada", "Tipo", "Necesidades", "Expectativas", "Seguimiento"],
        rows,
    )


def _render_caracterizacion(doc: Document, content: dict):
    items = [p for p in _as_list(content.get("characterizations")) if isinstance(p, dict)]
    if not items:
        _add_paragraph(doc, "Sin caracterizaciones registradas aún.", italic=True)
        return
    for proc in items:
        _add_heading(doc, _as_str(proc.get("process_name"), "Proceso"), level=2)
        _add_paragraph(doc, f"Objetivo: {_as_str(proc.get('objective'))}")
        _add_paragraph(doc, f"Alcance: {_as_str(proc.get('scope'))}")
        _add_paragraph(doc, f"Responsable: {_as_str(proc.get('owner'))}")
        for key, title in (
            ("inputs", "Entradas"),
            ("outputs", "Salidas"),
            ("main_activities", "Actividades principales"),
        ):
            values = _as_list(proc.get(key))
            if values:
                _add_heading(doc, title, level=3)
                _add_bullets(doc, values)


def _render_matriz(doc: Document, content: dict):
    rows = []
    for r in _as_list(content.get("interactions")):
        if not isinstance(r, dict):
            continue
        rows.append(
            [
                _as_str(r.get("source_process")),
                _as_str(r.get("target_process")),
                _as_str(r.get("information_transferred")),
                _as_str(r.get("dependency")),
            ]
        )
    _add_table(
        doc,
        ["Proceso origen", "Proceso destino", "Información transferida", "Dependencia"],
        rows,
    )


def _render_legal(doc: Document, content: dict):
    if content.get("summary"):
        _add_paragraph(doc, _as_str(content["summary"]))
    rows = []
    for r in _as_list(content.get("requirements")):
        if not isinstance(r, dict):
            continue
        rows.append(
            [
                _as_str(r.get("law_or_regulation")),
                _as_str(r.get("requirement_summary")),
                _as_str(r.get("compliance_status")),
                _as_str(r.get("evidence")),
                _as_str(r.get("responsible")),
            ]
        )
    _add_table(
        doc,
        ["Ley / norma", "Requisito", "Cumplimiento", "Evidencia", "Responsable"],
        rows,
    )


def _render_politica(doc: Document, content: dict):
    _add_paragraph(doc, _as_str(content.get("policy_text"), "Política en elaboración."))
    commitments = _as_list(content.get("commitments"))
    if commitments:
        _add_heading(doc, "Compromisos", level=2)
        _add_bullets(doc, commitments)
    if content.get("alignment_with_context"):
        _add_paragraph(doc, _as_str(content["alignment_with_context"]), italic=True)


def _render_objetivos(doc: Document, content: dict):
    rows = []
    for o in _as_list(content.get("objectives")):
        if not isinstance(o, dict):
            continue
        rows.append(
            [
                _as_str(o.get("objective")),
                _as_str(o.get("indicator")),
                _as_str(o.get("target")),
                _as_str(o.get("deadline")),
                _as_str(o.get("responsible")),
                _as_str(o.get("linked_process")),
            ]
        )
    _add_table(
        doc,
        ["Objetivo", "Indicador", "Meta", "Plazo", "Responsable", "Proceso"],
        rows,
    )
    if content.get("alignment_with_policy"):
        _add_heading(doc, "Alineación con la política", level=2)
        _add_paragraph(doc, _as_str(content["alignment_with_policy"]))


def _render_procedimientos(doc: Document, content: dict):
    procedures = [p for p in _as_list(content.get("procedures")) if isinstance(p, dict)]
    if not procedures:
        _add_paragraph(doc, "Sin procedimientos definidos aún.", italic=True)
        return
    for idx, proc in enumerate(procedures, start=1):
        code = _as_str(proc.get("code"), f"PROC-{idx}")
        title = _as_str(proc.get("title"), "Procedimiento")
        _add_heading(doc, f"{code} — {title}", level=2)
        _add_paragraph(doc, f"Proceso: {_as_str(proc.get('process_name'))}")
        _add_paragraph(doc, f"Alcance: {_as_str(proc.get('scope'))}")
        _add_paragraph(doc, f"Objetivo: {_as_str(proc.get('objective'))}")
        activities = [a for a in _as_list(proc.get("activities")) if isinstance(a, dict)]
        if activities:
            _add_heading(doc, "Actividades", level=3)
            for i, act in enumerate(activities, start=1):
                step = _as_str(act.get("step"), str(i))
                desc = _as_str(act.get("description"))
                responsible = _as_str(act.get("responsible"))
                line = f"{step}. {desc}"
                if responsible:
                    line += f" (Responsable: {responsible})"
                p = doc.add_paragraph(style="List Number")
                run = p.add_run(line)
                _set_run_font(run, size=11)


def _render_indicadores(doc: Document, content: dict):
    rows = []
    for ind in _as_list(content.get("indicators")):
        if not isinstance(ind, dict):
            continue
        rows.append(
            [
                _as_str(ind.get("process_name")),
                _as_str(ind.get("name")),
                _as_str(ind.get("objective")),
                _as_str(ind.get("formula")),
                _as_str(ind.get("frequency")),
                _as_str(ind.get("target")),
                _as_str(ind.get("responsible")),
                _as_str(ind.get("data_source")),
            ]
        )
    _add_table(
        doc,
        ["Proceso", "Indicador", "Objetivo", "Fórmula", "Frecuencia", "Meta", "Responsable", "Fuente"],
        rows,
    )


def _render_riesgos(doc: Document, content: dict):
    rows = []
    for e in _as_list(content.get("entries")):
        if not isinstance(e, dict):
            continue
        rows.append(
            [
                _as_str(e.get("related_process")),
                _as_str(e.get("risk")),
                _as_str(e.get("opportunity"), "—"),
                _as_str(e.get("risk_level")),
                _as_str(e.get("proposed_action")),
                _as_str(e.get("responsible")),
            ]
        )
    _add_table(
        doc,
        ["Proceso", "Riesgo", "Oportunidad", "Nivel", "Acción propuesta", "Responsable"],
        rows,
    )


def _render_registros(doc: Document, content: dict):
    if content.get("summary"):
        _add_paragraph(doc, _as_str(content["summary"]))
    rows = []
    for r in _as_list(content.get("records")):
        if not isinstance(r, dict):
            continue
        rows.append(
            [
                _as_str(r.get("code")),
                _as_str(r.get("name")),
                _as_str(r.get("related_clause")),
                _as_str(r.get("related_process")),
                _as_str(r.get("retention_period")),
                _as_str(r.get("responsible")),
            ]
        )
    _add_table(
        doc,
        ["Código", "Registro", "Cláusula", "Proceso", "Conservación", "Responsable"],
        rows,
    )


def _render_diagram_placeholder(doc: Document, component_type: str, content: dict):
    _add_paragraph(
        doc,
        "Este documento es gráfico en la plataforma. A continuación se incluye la información textual disponible para edición.",
        italic=True,
    )
    if component_type == "mapa_procesos":
        processes = _as_list(content.get("processes"))
        if processes:
            _add_heading(doc, "Procesos", level=2)
            rows = []
            for p in processes:
                if isinstance(p, dict):
                    rows.append(
                        [
                            _as_str(p.get("name") or p.get("process_name")),
                            _as_str(p.get("category") or p.get("type")),
                            _as_str(p.get("owner") or p.get("responsible")),
                            _as_str(p.get("description")),
                        ]
                    )
                else:
                    rows.append([_as_str(p), "", "", ""])
            _add_table(doc, ["Proceso", "Categoría", "Responsable", "Descripción"], rows)
        else:
            _render_generic(doc, content, skip_keys={"organization_name"})
    elif component_type == "diagrama_flujo":
        diagrams = [d for d in _as_list(content.get("diagrams")) if isinstance(d, dict)]
        if not diagrams:
            _render_generic(doc, content, skip_keys={"organization_name"})
            return
        for d in diagrams:
            _add_heading(doc, _as_str(d.get("process_name"), "Proceso"), level=2)
            steps = [s for s in _as_list(d.get("steps") or d.get("activities")) if isinstance(s, dict)]
            if steps:
                rows = [
                    [
                        _as_str(s.get("step") or s.get("id") or s.get("order")),
                        _as_str(s.get("name") or s.get("activity") or s.get("description")),
                        _as_str(s.get("role") or s.get("responsible")),
                        _as_str(s.get("type") or s.get("step_type")),
                    ]
                    for s in steps
                ]
                _add_table(doc, ["Paso", "Actividad", "Rol", "Tipo"], rows)
            else:
                _render_generic(doc, d, skip_keys={"process_name"})
    elif component_type == "organigrama":
        nodes = [n for n in _as_list(content.get("nodes")) if isinstance(n, dict)]
        if nodes:
            rows = [
                [
                    _as_str(n.get("name") or n.get("title") or n.get("role")),
                    _as_str(n.get("area") or n.get("department")),
                    _as_str(n.get("reports_to") or n.get("parent")),
                    _as_str(n.get("responsibilities") or n.get("description")),
                ]
                for n in nodes
            ]
            _add_table(doc, ["Cargo / persona", "Área", "Reporta a", "Descripción"], rows)
        else:
            _render_generic(doc, content, skip_keys={"organization_name"})


def _render_generic(doc: Document, content: dict, skip_keys: set[str] | None = None):
    skip = skip_keys or set()
    useful = {
        k: v
        for k, v in content.items()
        if k not in skip and v not in (None, "", [], {})
    }
    if not useful:
        _add_paragraph(doc, "Documento sin contenido textual aún.", italic=True)
        return

    for key, value in useful.items():
        _add_heading(doc, _label(key), level=2)
        if isinstance(value, str):
            _add_paragraph(doc, value)
        elif isinstance(value, list):
            if value and all(isinstance(i, dict) for i in value):
                headers = []
                for item in value:
                    for k in item.keys():
                        if k not in headers:
                            headers.append(k)
                rows = [
                    [_as_str(item.get(h)) for h in headers]
                    for item in value
                ]
                _add_table(doc, [_label(h) for h in headers], rows)
            else:
                _add_bullets(doc, value)
        elif isinstance(value, dict):
            for sub_k, sub_v in value.items():
                _add_paragraph(doc, f"{_label(sub_k)}: {_as_str(sub_v)}")
        else:
            _add_paragraph(doc, _as_str(value))


RENDERERS = {
    "contexto_organizacion": _render_contexto,
    "alcance_sgc": _render_alcance,
    "partes_interesadas": _render_partes,
    "caracterizacion_procesos": _render_caracterizacion,
    "matriz_interaccion": _render_matriz,
    "cumplimiento_legal": _render_legal,
    "politica_calidad": _render_politica,
    "objetivos_calidad": _render_objetivos,
    "procedimientos": _render_procedimientos,
    "indicadores": _render_indicadores,
    "riesgos_oportunidades": _render_riesgos,
    "registros_requeridos": _render_registros,
}


def _ordered_docs(documents: dict[str, Any]) -> list[tuple[str, dict]]:
    ordered: list[tuple[str, dict]] = []
    seen: set[str] = set()
    for doc_type in PROGRESSIVE_DOC_TYPES:
        doc = documents.get(doc_type)
        if isinstance(doc, dict):
            ordered.append((doc_type, doc))
            seen.add(doc_type)
    for doc_type, doc in documents.items():
        if doc_type in seen or not isinstance(doc, dict):
            continue
        ordered.append((doc_type, doc))
    return ordered


def _has_meaningful_content(doc: dict) -> bool:
    content = doc.get("content")
    if isinstance(content, dict) and any(
        v not in (None, "", [], {}) for v in content.values()
    ):
        return True
    percent = doc.get("completeness_percent")
    try:
        if percent is not None and float(percent) > 0:
            return True
    except (TypeError, ValueError):
        pass
    status = _as_str(doc.get("status")).lower()
    return status in {"generated", "completed", "completo", "draft", "borrador"}


def _safe_filename(name: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]+', "", name)
    cleaned = re.sub(r"\s+", "_", cleaned.strip())
    return cleaned[:80] or "proyecto"


def build_consolidated_sgq_docx(
    *,
    project_name: str,
    organization_name: str,
    documents: dict[str, Any],
) -> tuple[bytes, str, int]:
    """
    Genera un .docx con todos los documentos SGQ ordenados.
    Retorna (bytes, filename, documentos_incluidos).
    """
    ordered = [
        (doc_type, doc)
        for doc_type, doc in _ordered_docs(documents)
        if _has_meaningful_content(doc)
    ]
    if not ordered:
        raise ValueError(
            "No hay documentos SGC con contenido para exportar. "
            "Complete la entrevista o genere borradores en la pestaña Documentos."
        )

    docx = Document()
    section = docx.sections[0]
    section.top_margin = Pt(56)
    section.bottom_margin = Pt(56)
    section.left_margin = Pt(56)
    section.right_margin = Pt(56)

    title = docx.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Documentación consolidada del SGC")
    _set_run_font(run, bold=True, size=20)

    org = docx.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org.add_run(organization_name or "Organización")
    _set_run_font(run, bold=True, size=14, color=RGBColor(0x1E, 0x3A, 0x5F))

    meta = docx.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    run = meta.add_run(
        f"Proyecto: {project_name}\n"
        f"Documentos incluidos: {len(ordered)}\n"
        f"Generado: {stamp}\n"
        "Archivo editable en Microsoft Word / Google Docs / LibreOffice"
    )
    _set_run_font(run, size=10, color=RGBColor(0x66, 0x66, 0x66))

    _add_paragraph(
        docx,
        "Este archivo reúne en un solo documento Word los borradores SGC del proyecto, "
        "en el orden oficial de la plataforma, para revisión y edición.",
        italic=True,
    )

    for index, (doc_type, record) in enumerate(ordered):
        docx.add_page_break()
        title_text = (
            DOCUMENT_TITLES.get(doc_type)
            or _as_str(record.get("title"))
            or _label(doc_type)
        )
        _add_heading(docx, f"{index + 1}. {title_text}", level=1)

        completeness = record.get("completeness_percent")
        status = _as_str(record.get("status"))
        info_bits = []
        if completeness is not None:
            info_bits.append(f"Completitud: {completeness}%")
        if status:
            info_bits.append(f"Estado: {status}")
        if info_bits:
            _add_paragraph(docx, " · ".join(info_bits), italic=True)

        content = record.get("content") if isinstance(record.get("content"), dict) else {}
        if doc_type in DIAGRAM_TYPES:
            _render_diagram_placeholder(docx, doc_type, content)
        elif doc_type in RENDERERS:
            RENDERERS[doc_type](docx, content)
        else:
            _render_generic(docx, content)

    buffer = io.BytesIO()
    docx.save(buffer)
    filename = (
        f"SGC_consolidado_{_safe_filename(organization_name or project_name)}.docx"
    )
    return buffer.getvalue(), filename, len(ordered)
